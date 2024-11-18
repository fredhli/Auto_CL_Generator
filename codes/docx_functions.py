import shutil
import os
import re

from ast import literal_eval
from docx import Document
from docx.shared import Inches
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from config import *
from gpt_functions import *


def cv_doc_to_text(cv_path, headers=headers):
    # Read the docx file
    doc = Document(cv_path)

    # Keep only the first page of the document
    first_page_elements = []
    for element in doc.element.body:
        first_page_elements.append(element)
        if element.tag.endswith("sectPr"):
            break
    doc.element.body.clear_content()
    for element in first_page_elements:
        doc.element.body.append(element)

    # Initialize text variable
    cv_text = ""

    # Iterate over elements in document body
    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = element.xpath(".//w:t")
            paragraph_text = "".join(node.text for node in para if node.text)

            if element.xpath(".//w:numPr"):  # Checking if paragraph is a list item
                paragraph_text = "- " + paragraph_text

            # Replace '\t' with ', ' and four spaces with ', ' if they exist
            paragraph_text = paragraph_text.replace("\t", ", ").replace("    ", ", ")
            cv_text += paragraph_text + "\n"

        elif isinstance(element, CT_Tbl):
            for table in doc.tables:
                for row in table.rows:
                    row_text = ": ".join(
                        cell.text.replace("\t", ", ").replace("    ", ", ")
                        if cell.text
                        else ""
                        for cell in row.cells
                    )
                    cv_text += row_text + "\n"

    for header in headers:
        cv_text = re.sub(rf"\b{header}\b", f"**{header}**", cv_text)
    return cv_text


def test_jd_keywords(jd, cv_text):
    system_msg = f"""You are a helpful assistant to detect whether my CV matches the job description in \
terms of keywords and strengths mentioned in the job description. Here is my cv in text\
format:\n\n{cv_text}\n\n Please check whether my CV lacks any keywords or strengths \
mentioned in the job description. Format your output to a python list. For example, \
['keyword1', 'keyword2', 'keyword3']. Do not include anything else in your response."""

    # Generate the response
    response = chatgpt(model="gpt-4o", prompt=jd, system_msg=system_msg, temp=0)

    return response


def get_cv_suggestion(kw_list, cv_text):
    system_msg = f"You are a helpful assistant to suggest special tailoring to my CV, catering to \
keywords and strengths mentioned in the job description. Here are the keywords and strengths mentioned\
in the job description: {kw_list}, I will attach my CV later. Please suggest how any bullet points \
in my CV that starts with '- ' can be modified to include any of the keywords or strengths mentioned \
in the job description. Format your output to a python dictionary. For example, \
{{'original bullet point 1': 'modified bullet point 1', 'original bullet point 2': 'modified bullet point 2'}}. \
Do not include anything else in your response."

    # retry = 0
    # max_retry = 5
    # while retry < max_retry:
    response = chatgpt(model="gpt-4o", prompt=cv_text, system_msg=system_msg, temp=0)

    response = response.replace("```python", "").replace("```", "")
    response = response.strip()

    #     try:
    #         response_dict = literal_eval(response)
    #         if isinstance(response_dict, dict):
    #             return response_dict
    #     except:
    #         pass
    #     retry += 1
    # return {}
    return literal_eval(response)


def change_cv_accordingly(company_name, change_dict, cv_path):
    # Copy the original CV to a new CV
    new_cv_path = cv_path.replace(".docx", f"_{company_name}_tailored.docx")
    shutil.copy(cv_path, new_cv_path)

    # Read the copied docx file
    doc = Document(new_cv_path)

    for old, new in change_dict.items():
        old = re.sub(r"^- ", "", old).strip()
        new = re.sub(r"^- ", "", new).strip()
        len_old = len(old.split())
        suggested_len_new = len(new.split())

        if suggested_len_new > len_old + 3:
            # Ask ChatGPT to shorten the bullet point until it is less than len_old + 3 words
            retry = 0
            max_retry = 5
            while retry < max_retry:
                prompt = f"Please help me rephrase this bullet point to strictly less than {len_old} words: {new}. Most importantly, make sure the word count does not exceed. Moreover, here is my old version: {old}.If you feel that your shortened version is not as good as the original, please just output the original. Thank you!"
                response = chatgpt(model="gpt-4o", prompt=prompt, temp=0.1 * retry)
                if len(response.split()) <= len_old:
                    new = response
                    break
                retry += 1

        # Direct Replace
        for para in doc.paragraphs:
            if para.text.strip() == old.strip():
                # Clear all runs within the paragraph
                for run in para.runs:
                    run.clear()  # Clear each run to remove the old text
                # Add the new text as a single run
                para.add_run(new)

        print(f"Replaced '{old}' with '{new}'")

    doc.save(new_cv_path)
    return new_cv_path
